"""Local document parsing and chunking for supported file types.

Supports: .txt, .md, .json, .csv, .pdf, .docx, .xlsx

All parsing is local. No cloud services, no OCR.
PDF: text extraction only (pypdf). Scanned image PDFs will show an error.
DOCX: paragraphs, headings, tables (python-docx).
XLSX: sheet detection, profiling, searchable text (openpyxl).
CSV: profiling with column type detection, missing value analysis.
"""

from __future__ import annotations

import csv
import hashlib
import io
import json as json_lib
import re
from abc import ABC, abstractmethod
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


def _make_chunk_id(source_id: str, chunk_index: int, text: str) -> str:
    """Generate a deterministic chunk ID from source_id, index, and text content.

    Re-parsing the same content produces the same ID, keeping evidence
    references stable across parse/index operations.
    """
    raw = f"{source_id}:{chunk_index}:{text.strip()[:200]}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16]


from decision_system.data_sources.models import (
    DataSourceChunk,
    ParsedBlock,
    ParseResult,
)

# ---------------------------------------------------------------------------
# Parser base class
# ---------------------------------------------------------------------------


class BaseParser(ABC):
    """Abstract base for all document parsers."""

    name: str = "base"
    version: str | None = None
    supported_extensions: list[str] = []
    supported_mime_types: list[str] = []

    @abstractmethod
    def parse(self, path: Path, source_id: str, workspace_id: str) -> ParseResult:
        """Parse a file and return a ParseResult."""
        ...

    def chunk(
        self,
        text: str,
        source_id: str,
        workspace_id: str,
        source_metadata: dict[str, Any] | None = None,
        page_number: int | None = None,
        sheet_name: str | None = None,
        block_type: str | None = None,
    ) -> list[DataSourceChunk]:
        """Split text into chunks with source metadata preserved."""
        if not text or not text.strip():
            return []

        paragraphs = re.split(r"\n\s*\n", text)
        chunks: list[DataSourceChunk] = []
        current_text = ""
        chunk_index = 0
        meta = dict(source_metadata or {})
        if page_number is not None:
            meta["page_number"] = page_number
        if sheet_name is not None:
            meta["sheet_name"] = sheet_name
        if block_type is not None:
            meta["block_type"] = block_type
        meta["parser"] = self.name

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(current_text) + len(para) < 1000:
                if current_text:
                    current_text += "\n\n" + para
                else:
                    current_text = para
            else:
                if current_text:
                    chunks.append(
                        DataSourceChunk(
                            chunk_id=_make_chunk_id(source_id, chunk_index, current_text),
                            source_id=source_id,
                            workspace_id=workspace_id,
                            chunk_index=chunk_index,
                            text=current_text,
                            metadata=dict(meta),
                            created_at=datetime.now(timezone.utc),
                        )
                    )
                    chunk_index += 1
                current_text = para

        if current_text:
            chunks.append(
                DataSourceChunk(
                    chunk_id=_make_chunk_id(source_id, chunk_index, current_text),
                    source_id=source_id,
                    workspace_id=workspace_id,
                    chunk_index=chunk_index,
                    text=current_text,
                    metadata=dict(meta),
                    created_at=datetime.now(timezone.utc),
                )
            )

        return chunks


# ---------------------------------------------------------------------------
# Text / Markdown parser
# ---------------------------------------------------------------------------


class TextParser(BaseParser):
    name = "text"
    supported_extensions = [".txt", ".md"]
    supported_mime_types = ["text/plain", "text/markdown"]

    def parse(self, path: Path, source_id: str, workspace_id: str) -> ParseResult:
        text = path.read_text(encoding="utf-8", errors="replace")
        if not text.strip():
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=["File is empty"],
                parser_name=self.name,
            )
        chunks = self.chunk(
            text, source_id, workspace_id, source_metadata={"source_name": path.name}
        )
        return ParseResult(
            source_id=source_id,
            workspace_id=workspace_id,
            text=text,
            metadata={"char_count": len(text), "line_count": text.count("\n") + 1},
            parser_name=self.name,
            chunks=chunks,
        )


# ---------------------------------------------------------------------------
# JSON parser
# ---------------------------------------------------------------------------


class JsonParser(BaseParser):
    name = "json"
    supported_extensions = [".json"]
    supported_mime_types = ["application/json"]

    def parse(self, path: Path, source_id: str, workspace_id: str) -> ParseResult:
        text = path.read_text(encoding="utf-8", errors="replace")
        if not text.strip():
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=["File is empty"],
                parser_name=self.name,
            )
        try:
            data = json_lib.loads(text)
        except json_lib.JSONDecodeError as e:
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=[f"JSON parsing error: {e}"],
                parser_name=self.name,
            )

        chunks: list[DataSourceChunk] = []
        if isinstance(data, dict):
            for key, value in data.items():
                chunk_text = f"{key}: {json_lib.dumps(value, indent=2, ensure_ascii=False)}"
                chunks.extend(
                    self.chunk(
                        chunk_text,
                        source_id,
                        workspace_id,
                        source_metadata={"source_name": path.name, "key": key},
                    )
                )
        elif isinstance(data, list):
            for i, item in enumerate(data):
                chunk_text = json_lib.dumps(item, indent=2, ensure_ascii=False)
                chunks.extend(
                    self.chunk(
                        chunk_text,
                        source_id,
                        workspace_id,
                        source_metadata={"source_name": path.name, "index": i},
                    )
                )
        else:
            chunk_text = str(data)
            chunks = self.chunk(
                chunk_text,
                source_id,
                workspace_id,
                source_metadata={"source_name": path.name},
            )

        return ParseResult(
            source_id=source_id,
            workspace_id=workspace_id,
            text=text,
            parser_name=self.name,
            chunks=chunks,
            metadata={"top_level_type": type(data).__name__},
        )


# ---------------------------------------------------------------------------
# PDF parser (text-based via pypdf)
# ---------------------------------------------------------------------------


class PdfParser(BaseParser):
    name = "pdf"
    version = "pypdf"
    supported_extensions = [".pdf"]
    supported_mime_types = ["application/pdf"]

    def parse(self, path: Path, source_id: str, workspace_id: str) -> ParseResult:
        try:
            from pypdf import PdfReader
        except ImportError:
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=["pypdf is not installed. Install it with: pip install pypdf"],
                parser_name=self.name,
            )
        warnings: list[str] = []
        try:
            reader = PdfReader(path)
        except Exception as e:
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=[f"Failed to open PDF: {e}"],
                parser_name=self.name,
            )

        if reader.is_encrypted:
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=["PDF is encrypted and cannot be read by this local parser."],
                parser_name=self.name,
            )

        pages_parsed: list[dict[str, Any]] = []
        full_text_parts: list[str] = []
        all_chunks: list[DataSourceChunk] = []
        page_count = len(reader.pages)

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            char_count = len(text)
            if not text.strip():
                warnings.append(f"Page {page_num + 1} has no extractable text")
                continue
            pages_parsed.append(
                {
                    "page_number": page_num + 1,
                    "text": text,
                    "char_count": char_count,
                }
            )
            full_text_parts.append(text)
            chunks = self.chunk(
                text,
                source_id,
                workspace_id,
                source_metadata={"source_name": path.name},
                page_number=page_num + 1,
            )
            all_chunks.extend(chunks)

        full_text = "\n\n".join(full_text_parts)
        metadata: dict[str, Any] = {
            "page_count": page_count,
        }
        if reader.metadata:
            if reader.metadata.title:
                metadata["title"] = reader.metadata.title
            if reader.metadata.author:
                metadata["author"] = reader.metadata.author

        if not full_text.strip():
            warnings.append(
                "PDF contains no extractable text. OCR is not supported in this local parser."
            )

        return ParseResult(
            source_id=source_id,
            workspace_id=workspace_id,
            text=full_text,
            pages=pages_parsed,
            metadata=metadata,
            warnings=warnings,
            parser_name=self.name,
            parser_version="pypdf",
            chunks=all_chunks,
        )


# ---------------------------------------------------------------------------
# DOCX parser (via python-docx)
# ---------------------------------------------------------------------------


class DocxParser(BaseParser):
    name = "docx"
    version = "python-docx"
    supported_extensions = [".docx"]
    supported_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def parse(self, path: Path, source_id: str, workspace_id: str) -> ParseResult:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=["python-docx is not installed. Install it with: pip install python-docx"],
                parser_name=self.name,
            )
        warnings: list[str] = []
        try:
            doc = DocxDocument(str(path))
        except Exception as e:
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=[f"Failed to open DOCX: {e}"],
                parser_name=self.name,
            )

        blocks: list[ParsedBlock] = []
        all_chunks: list[DataSourceChunk] = []
        full_text_parts: list[str] = []
        para_count = 0
        table_count = 0

        for block_idx, element in enumerate(doc.element.body):
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                para = doc.paragraphs[para_count] if para_count < len(doc.paragraphs) else None
                if para is None:
                    para_count += 1
                    continue
                text = para.text.strip()
                if not text:
                    para_count += 1
                    continue
                style = para.style.name if para.style else ""
                is_heading = style.startswith("Heading") if style else False
                block_type = "heading" if is_heading else "paragraph"
                blocks.append(
                    ParsedBlock(
                        block_type=block_type,
                        text=text,
                        index=block_idx,
                        metadata={"style": style},
                    )
                )
                full_text_parts.append(text)
                chunks = self.chunk(
                    text,
                    source_id,
                    workspace_id,
                    source_metadata={"source_name": path.name},
                    block_type=block_type,
                )
                all_chunks.extend(chunks)
                para_count += 1

            elif tag == "tbl":
                table = doc.tables[table_count] if table_count < len(doc.tables) else None
                if table is None:
                    table_count += 1
                    continue
                rows_text: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_text.append(" | ".join(cells))
                table_text = "\n".join(rows_text)
                blocks.append(
                    ParsedBlock(
                        block_type="table",
                        text=table_text,
                        index=block_idx,
                        metadata={"rows": len(table.rows), "cols": len(table.columns)},
                    )
                )
                full_text_parts.append(f"[TABLE]:\n{table_text}")
                chunks = self.chunk(
                    table_text,
                    source_id,
                    workspace_id,
                    source_metadata={"source_name": path.name},
                    block_type="table",
                )
                all_chunks.extend(chunks)
                table_count += 1

        full_text = "\n\n".join(full_text_parts)

        if not full_text.strip():
            warnings.append("DOCX appears to have no extractable text content")

        metadata = {
            "paragraph_count": para_count,
            "table_count": table_count,
        }

        return ParseResult(
            source_id=source_id,
            workspace_id=workspace_id,
            text=full_text,
            tables=[b.model_dump() for b in blocks if b.block_type == "table"],
            metadata=metadata,
            warnings=warnings,
            parser_name=self.name,
            parser_version="python-docx",
            chunks=all_chunks,
        )


# ---------------------------------------------------------------------------
# XLSX parser and profiler (via openpyxl)
# ---------------------------------------------------------------------------


class XlsxParser(BaseParser):
    name = "xlsx"
    version = "openpyxl"
    supported_extensions = [".xlsx"]
    supported_mime_types = ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]

    def parse(self, path: Path, source_id: str, workspace_id: str) -> ParseResult:
        try:
            import openpyxl
        except ImportError:
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=["openpyxl is not installed. Install with: pip install openpyxl"],
                parser_name=self.name,
            )
        warnings: list[str] = []
        try:
            wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        except Exception as e:
            return ParseResult(
                source_id=source_id,
                workspace_id=workspace_id,
                warnings=[f"Failed to open XLSX: {e}"],
                parser_name=self.name,
            )

        sheets_parsed: list[dict[str, Any]] = []
        full_text_parts: list[str] = []
        all_chunks: list[DataSourceChunk] = []
        sheets = wb.sheetnames
        sheet_count = len(sheets)

        for sheet_name in sheets:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                warnings.append(f"Sheet '{sheet_name}' is empty")
                continue

            row_count = len(rows)
            col_count = max((len(r) for r in rows), default=0)
            headers = [str(c) if c is not None else "" for c in (rows[0] if rows else [])]

            # Build searchable text representation
            sheet_text = f"Sheet: {sheet_name}\n"
            if headers:
                sheet_text += "Columns: " + ", ".join(headers) + "\n"
            sheet_text += f"Rows: {row_count}\n\n"

            sample_rows = rows[:5]
            for i, row in enumerate(sample_rows):
                vals = [str(c) if c is not None else "" for c in row]
                sheet_text += f"Row {i + 1}: " + " | ".join(vals) + "\n"

            full_text_parts.append(sheet_text)
            sheets_parsed.append(
                {
                    "sheet_name": sheet_name,
                    "row_count": row_count,
                    "column_count": col_count,
                    "headers": headers,
                }
            )

            chunks = self.chunk(
                sheet_text,
                source_id,
                workspace_id,
                source_metadata={"source_name": path.name},
                sheet_name=sheet_name,
            )
            all_chunks.extend(chunks)

            # Also create chunks per-row for finer granularity
            for i, row in enumerate(rows[1:], start=2):
                vals = [str(c) if c is not None else "" for c in row]
                row_text = " | ".join(vals)
                if row_text.strip():
                    row_chunks = self.chunk(
                        f"{sheet_name} Row {i}: {row_text}",
                        source_id,
                        workspace_id,
                        source_metadata={"source_name": path.name},
                        sheet_name=sheet_name,
                    )
                    all_chunks.extend(row_chunks)

        wb.close()
        full_text = "\n\n".join(full_text_parts)

        metadata = {
            "sheet_count": sheet_count,
            "sheet_names": sheets,
        }

        return ParseResult(
            source_id=source_id,
            workspace_id=workspace_id,
            text=full_text,
            tables=sheets_parsed,
            metadata=metadata,
            warnings=warnings,
            parser_name=self.name,
            parser_version="openpyxl",
            chunks=all_chunks,
        )

    def profile(self, path: Path, source_id: str, workspace_id: str) -> dict[str, Any]:
        """Generate a dataset profile for an XLSX workbook."""
        try:
            import openpyxl
        except ImportError:
            return {"error": "openpyxl not installed", "source_id": source_id}
        try:
            wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        except Exception as e:
            return {"error": str(e), "source_id": source_id}

        sheets_profile: list[dict[str, Any]] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            row_count = len(rows)
            col_count = max((len(r) for r in rows), default=0) if rows else 0
            headers = [str(c) if c is not None else "" for c in (rows[0] if rows else [])]
            missing_counts: dict[str, int] = {}
            numeric_cols: dict[str, list[float]] = {}
            categorical_cols: dict[str, dict[str, int]] = {}

            for h in headers:
                missing_counts[h] = 0
                numeric_cols[h] = []
                categorical_cols[h] = {}

            for row in rows[1:]:
                for idx, h in enumerate(headers):
                    val = row[idx] if idx < len(row) else None
                    if val is None or (isinstance(val, str) and not val.strip()):
                        missing_counts[h] += 1
                        continue
                    try:
                        numeric_cols[h].append(float(val))
                    except (ValueError, TypeError):
                        cat = str(val)
                        categorical_cols[h][cat] = categorical_cols[h].get(cat, 0) + 1

            col_profiles = []
            for h in headers:
                col_info = {
                    "name": h,
                    "missing": missing_counts.get(h, 0),
                    "missing_pct": round(missing_counts.get(h, 0) / max(row_count - 1, 1), 4)
                    if row_count > 1
                    else 0,
                }
                if numeric_cols.get(h):
                    nums = numeric_cols[h]
                    col_info["dtype"] = "numeric"
                    col_info["min"] = min(nums)
                    col_info["max"] = max(nums)
                    col_info["mean"] = round(sum(nums) / len(nums), 4)
                    col_info["count"] = len(nums)
                else:
                    col_info["dtype"] = "categorical"
                    top_cats = sorted(categorical_cols[h].items(), key=lambda x: -x[1])[:10]
                    col_info["top_values"] = top_cats
                col_profiles.append(col_info)

            warnings_local: list[str] = []
            for cp in col_profiles:
                if cp.get("missing_pct", 0) > 0.5:
                    warnings_local.append(f"Column '{cp['name']}' is >50% missing")

            sheets_profile.append(
                {
                    "sheet_name": sheet_name,
                    "row_count": row_count,
                    "column_count": col_count,
                    "columns": col_profiles,
                    "sample_rows": [
                        {
                            headers[j]: (
                                str(rows[i][j])
                                if j < len(rows[i]) and rows[i][j] is not None
                                else ""
                            )
                            for j in range(len(headers))
                        }
                        for i in range(1, min(6, row_count))
                    ]
                    if row_count > 1
                    else [],
                    "warnings": warnings_local,
                }
            )

        sheet_names = (
            list(wb.sheetnames)
            if hasattr(wb, "sheetnames")
            else [s["sheet_name"] for s in sheets_profile]
        )
        sheet_count = len(sheet_names)
        wb.close()
        return {
            "source_id": source_id,
            "workspace_id": workspace_id,
            "sheet_count": sheet_count,
            "sheet_names": sheet_names,
            "sheets": sheets_profile,
        }


# ---------------------------------------------------------------------------
# Parser registry
# ---------------------------------------------------------------------------

_registry: dict[str, BaseParser] = {}


def _init_registry():
    """Initialize the parser registry with all available parsers."""
    parsers: list[BaseParser] = [
        TextParser(),
        JsonParser(),
        PdfParser(),
        DocxParser(),
        XlsxParser(),
    ]
    # Conditionally add OCR parsers for images
    try:
        from decision_system.data_sources.ocr_parser import ImageOcrParser

        for ext in ImageOcrParser.supported_extensions:
            _registry[ext.lower()] = ImageOcrParser()
    except ImportError:
        pass
    for parser in parsers:
        for ext in parser.supported_extensions:
            _registry[ext.lower()] = parser


def get_parser(ext: str) -> BaseParser | None:
    """Get the parser for a given file extension."""
    if not _registry:
        _init_registry()
    return _registry.get(ext.lower())


def get_supported_extensions() -> list[str]:
    """Return the list of all supported file extensions."""
    if not _registry:
        _init_registry()
    return list(_registry.keys())


def is_parsable(ext: str) -> bool:
    """Check if a file extension is supported."""
    if not _registry:
        _init_registry()
    return ext.lower() in _registry


def parse_document(
    path_or_content: str | Path, ext: str, source_id: str, workspace_id: str
) -> tuple[list[DataSourceChunk], list[str], dict[str, Any]]:
    """Parse a document and return (chunks, warnings, metadata).

    Accepts either a Path to a file on disk, or a string of content.
    If a string is passed, it is written to a temp file for parsing.
    """
    if not _registry:
        _init_registry()
    result = ParseResult(
        source_id=source_id,
        workspace_id=workspace_id,
        warnings=[],
        parser_name="unknown",
        chunks=[],
    )

    # Resolve path
    if isinstance(path_or_content, str):
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=ext, mode="w", delete=False, encoding="utf-8") as f:
            f.write(path_or_content)
            path = Path(f.name)
        try:
            chunks, warnings, metadata = _do_parse(path, ext, source_id, workspace_id, result)
            return chunks, warnings
        finally:
            path.unlink(missing_ok=True)
    else:
        chunks, warnings, metadata = _do_parse(
            path_or_content, ext, source_id, workspace_id, result
        )
        return chunks, warnings


def _do_parse(
    path: Path,
    ext: str,
    source_id: str,
    workspace_id: str,
    result: ParseResult,
) -> tuple[list[DataSourceChunk], list[str], dict[str, Any]]:
    """Internal dispatch to the appropriate parser."""
    if not path.exists():
        result.warnings.append(f"File not found: {path}")
        return result.chunks, result.warnings, result.metadata

    ext_lower = ext.lower()
    parser = _registry.get(ext_lower)

    if parser is None:
        result.warnings.append(f"Unsupported file extension: {ext}")
        return result.chunks, result.warnings, result.metadata

    try:
        result = parser.parse(path, source_id, workspace_id)
    except Exception as e:
        result.warnings.append(f"Parsing error: {e}")

    # OCR fallback: if PDF produced no text, try scanned PDF OCR
    if ext_lower == ".pdf" and result and not result.text.strip():
        try:
            from decision_system.data_sources.ocr_parser import ScannedPdfParser

            ocr_parser = ScannedPdfParser()
            ocr_result = ocr_parser.parse(path, source_id, workspace_id)
            if ocr_result and ocr_result.text.strip():
                result = ocr_result
                result.warnings.append("Text extracted via OCR (scanned PDF)")
        except ImportError:
            pass
        except Exception as e:
            result.warnings.append(f"OCR fallback attempted but failed: {e}")

    return result.chunks, result.warnings, result.metadata


def parse_document_from_content(
    content: str, ext: str, source_id: str, workspace_id: str
) -> tuple[list[DataSourceChunk], list[str]]:
    """Parse from string content (used for in-memory / test scenarios).

    Writes content to a temp file and dispatches to the right parser.
    """
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=ext, mode="w", delete=False, encoding="utf-8") as f:
        f.write(content)
        tmp_path = Path(f.name)
    try:
        chunks, warnings, metadata = parse_document(tmp_path, ext, source_id, workspace_id)
        return chunks, warnings
    finally:
        tmp_path.unlink(missing_ok=True)


# Legacy compatibility wrapper
def parse_text(content: str, source_id: str, workspace_id: str) -> list[DataSourceChunk]:
    """Legacy wrapper."""
    parser = TextParser()
    import tempfile
    from pathlib import Path

    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False, encoding="utf-8") as f:
        f.write(content)
        tmp_path = Path(f.name)
    try:
        result = parser.parse(tmp_path, source_id, workspace_id)
        return result.chunks
    finally:
        tmp_path.unlink(missing_ok=True)


def parse_json(content: str, source_id: str, workspace_id: str) -> list[DataSourceChunk]:
    """Legacy wrapper."""
    parser = JsonParser()
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".json", mode="w", delete=False, encoding="utf-8") as f:
        f.write(content)
        tmp_path = Path(f.name)
    try:
        result = parser.parse(tmp_path, source_id, workspace_id)
        return result.chunks
    finally:
        tmp_path.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# CSV profiling (unchanged from previous implementation)
# ---------------------------------------------------------------------------


def profile_csv(content: str, source_id: str, workspace_id: str) -> dict[str, Any]:
    """Profile CSV content and return a structured profile dict.

    Returns profile with row_count, column_count, column_types,
    missing_values, numeric_summary, categorical_summary, sample_rows.
    """
    try:
        reader = csv.DictReader(io.StringIO(content))
        rows = list(reader)
    except Exception as e:
        return {"error": str(e), "row_count": 0}

    if not rows:
        return {
            "source_id": source_id,
            "workspace_id": workspace_id,
            "row_count": 0,
            "column_count": 0,
            "columns": [],
            "warnings": ["No data rows found"],
        }

    headers = reader.fieldnames or []
    col_count = len(headers)
    row_count = len(rows)

    columns = []
    column_types: dict[str, str] = {}
    missing_values: dict[str, int] = {}
    numeric_summary: dict[str, dict[str, float]] = {}
    categorical_summary: dict[str, list[tuple[str, int]]] = {}
    date_like_columns: list[str] = []

    for header in headers:
        col_values = [r.get(header, "") or "" for r in rows]
        non_empty = [v for v in col_values if v.strip()]
        missing = row_count - len(non_empty)

        # Detect type
        is_numeric = False
        nums = []
        for v in non_empty:
            cleaned = v.strip().replace(",", "").replace("$", "").replace("%", "")
            try:
                nums.append(float(cleaned))
            except ValueError:
                pass

        if nums and len(nums) == len(non_empty):
            is_numeric = True
            column_types[header] = "numeric"
            numeric_summary[header] = {
                "min": min(nums),
                "max": max(nums),
                "mean": round(sum(nums) / len(nums), 4),
                "count": len(nums),
            }
        else:
            column_types[header] = "categorical"
            freqs: dict[str, int] = {}
            for v in non_empty:
                freqs[v] = freqs.get(v, 0) + 1
            categorical_summary[header] = sorted(freqs.items(), key=lambda x: (-x[1], x[0]))[:10]

        missing_values[header] = missing
        columns.append(
            {
                "name": header,
                "dtype": column_types[header],
                "missing": missing,
                "missing_pct": round(missing / row_count, 4) if row_count else 0,
            }
        )

        # Date-like detection
        date_pattern = re.compile(r"date|time|month|year|day|period|week", re.IGNORECASE)
        if date_pattern.search(header) and not is_numeric:
            date_like_columns.append(header)

    warnings = []
    for col in columns:
        if col["missing_pct"] > 0.5:
            warnings.append(f"Column '{col['name']}' is >50% missing")

    return {
        "source_id": source_id,
        "workspace_id": workspace_id,
        "row_count": row_count,
        "column_count": col_count,
        "columns": columns,
        "column_types": column_types,
        "missing_values": missing_values,
        "numeric_summary": numeric_summary,
        "categorical_summary": categorical_summary,
        "date_like_columns": date_like_columns,
        "sample_rows": rows[:5],
        "warnings": warnings,
    }


def profile_json_content(content: str, source_id: str, workspace_id: str) -> dict[str, Any]:
    """Profile JSON content and return a structured profile dict."""
    try:
        data = json_lib.loads(content)
    except json_lib.JSONDecodeError as e:
        return {
            "source_id": source_id,
            "workspace_id": workspace_id,
            "error": str(e),
            "top_level_type": "unknown",
        }

    result: dict[str, Any] = {
        "source_id": source_id,
        "workspace_id": workspace_id,
        "top_level_type": type(data).__name__,
    }

    if isinstance(data, list):
        result["record_count"] = len(data)
        result["sample_records"] = data[:3] if data else []
        if data and isinstance(data[0], dict):
            result["field_paths"] = sorted(data[0].keys())
    elif isinstance(data, dict):
        result["field_paths"] = sorted(data.keys())
        result["sample_records"] = [dict(list(data.items())[:5])]
    else:
        result["sample_records"] = [str(data)]

    return result
